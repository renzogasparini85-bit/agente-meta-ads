from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

VIOLET = RGBColor(0x6B, 0x21, 0xA8)
ORANGE = RGBColor(0xFF, 0x6B, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK = RGBColor(0x1F, 0x1F, 0x1F)
GRAY = RGBColor(0x6B, 0x72, 0x80)
LIGHT_BG = RGBColor(0xF5, 0xF3, 0xFF)

doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def set_cell_bg(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_heading(doc, text, level=1, color=VIOLET):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    if level == 1:
        run.font.size = Pt(20)
    elif level == 2:
        run.font.size = Pt(14)
    else:
        run.font.size = Pt(12)
    return p

def add_body(doc, text, bold=False, italic=False, color=DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    return p

def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.name = 'Calibri'
        r1.font.size = Pt(11)
        r1.font.color.rgb = DARK
        r2 = p.add_run(text)
        r2.font.name = 'Calibri'
        r2.font.size = Pt(11)
        r2.font.color.rgb = DARK
    else:
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.color.rgb = DARK

def add_table(doc, headers, rows, col_widths=None, header_bg='6B21A8'):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

    # Rows
    for ri, row in enumerate(rows):
        bg = 'F5F3FF' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = 'Calibri'
            run.font.size = Pt(10)
            run.font.color.rgb = DARK

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

# ── PORTADA ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(40)
p.paragraph_format.space_after = Pt(4)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SISTEMA DE INTELIGENCIA')
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = VIOLET
r.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(2)
r2 = p2.add_run('PUBLICITARIA CON IA')
r2.bold = True
r2.font.size = Pt(28)
r2.font.color.rgb = VIOLET
r2.font.name = 'Calibri'

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(16)
p3.paragraph_format.space_after = Pt(8)
r3 = p3.add_run('Propuesta para IFPA — Mayo 2026')
r3.font.size = Pt(14)
r3.font.color.rgb = GRAY
r3.font.name = 'Calibri'

doc.add_paragraph()

# Línea decorativa
p_line = doc.add_paragraph()
p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_line = p_line.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
r_line.font.color.rgb = ORANGE
r_line.font.size = Pt(14)

doc.add_paragraph()

# ── SECCIÓN 1 ─────────────────────────────────────────────────────────────────
add_heading(doc, '01. El problema', 1)
add_body(doc, 'Tu equipo de campañas trabaja en el Ads Manager, revisa números manualmente y toma decisiones con datos incompletos. Siempre llega tarde: el anuncio ya quemó presupuesto antes de que alguien lo pausara, el creativo ganador tardó dos semanas en escalarse.')
add_body(doc, 'Esto no es un problema de personas. Es un problema de velocidad y volumen de datos.')

# ── SECCIÓN 2 ─────────────────────────────────────────────────────────────────
add_heading(doc, '02. Lo que hace el sistema', 1)
add_body(doc, 'Un agente de inteligencia artificial que monitorea las campañas de Meta Ads en tiempo real, detecta oportunidades y problemas antes de que impacten en el presupuesto, y le entrega al equipo decisiones listas para ejecutar — no datos crudos.')

add_heading(doc, 'Todos los días', 2)
add_bullet(doc, 'Reporte automático en Telegram con estado de cada campaña: gasto, conversaciones, CPA, frecuencia y alertas', '📊 ')
add_bullet(doc, 'Detección de anomalías en tiempo real: anuncio gastando sin convertir, frecuencia alta, CTR en caída', '🚨 ')
add_bullet(doc, 'Ranking de creativos: cuáles están ganando, cuáles están muriendo, cuáles hay que replicar', '🏆 ')

add_heading(doc, 'Cada semana', 2)
add_bullet(doc, 'Auditoría completa con recomendación concreta por anuncio: escalar, pausar, replicar o renovar', '📋 ')
add_bullet(doc, 'Detección de joyas ocultas: creativos con bajo presupuesto y CPA excepcional que el algoritmo ignora', '💎 ')
add_bullet(doc, 'Comparativa de formatos: imagen vs carrusel vs video por campaña y sede', '📈 ')

add_heading(doc, 'Cada mes', 2)
add_bullet(doc, 'Brief de creativos con copies listos para el diseñador, basados en los ángulos que mejor convierten', '✍️ ')
add_bullet(doc, 'Análisis de audiencias: qué segmentos generan conversaciones de mejor calidad', '🎯 ')
add_bullet(doc, 'Proyección de presupuesto: dónde asignar más inversión para el mes siguiente', '💰 ')

add_heading(doc, 'Por comando desde WhatsApp o Telegram', 2)
add_bullet(doc, 'Pausar un anuncio')
add_bullet(doc, 'Escalar el presupuesto de un conjunto')
add_bullet(doc, 'Consultar el estado de una campaña')
add_bullet(doc, 'Pedir un reporte en cualquier momento')
add_body(doc, 'Sin entrar al Ads Manager.', bold=True, color=VIOLET)

# ── SECCIÓN 3 ─────────────────────────────────────────────────────────────────
doc.add_page_break()
add_heading(doc, '03. Resultados reales — Primera auditoría IFPA', 1)
add_body(doc, 'El sistema ya corre sobre las campañas de IFPA. Estos son los hallazgos de la primera semana:')
doc.add_paragraph()

add_table(doc,
    ['Hallazgo', 'Impacto estimado'],
    [
        ['Anuncios con $70.000+ semanal sin conversión → pausados', '~$280.000/mes recuperado'],
        ['Creativo con CPA $40 ignorado (promedio cuenta: $204)', 'Potencial 5x con mismo presupuesto'],
        ['Campaña Inglés con doble CPA por objetivo equivocado', '~$150.000/mes corregibles'],
        ['3 creativos ganadores identificados para replicar en otras sedes', 'Expansión sin costo de testeo'],
    ],
    col_widths=[10, 5.5]
)
doc.add_paragraph()
p_efic = doc.add_paragraph()
p_efic.paragraph_format.space_before = Pt(8)
r_efic = p_efic.add_run('Eficiencia recuperable estimada: 15–20% del presupuesto mensual.')
r_efic.bold = True
r_efic.font.size = Pt(12)
r_efic.font.color.rgb = VIOLET
r_efic.font.name = 'Calibri'

add_body(doc, 'Con una inversión de $3.000.000 ARS/mes en ads, eso representa entre $450.000 y $600.000 ARS/mes en valor recuperado.')

# ── SECCIÓN 4 ─────────────────────────────────────────────────────────────────
add_heading(doc, '04. Inversión', 1)
doc.add_paragraph()

add_table(doc,
    ['Plan', 'Precio mensual', 'Incluye'],
    [
        ['INTELIGENCIA', 'ARS 350.000', 'Reporte diario + alertas + auditoría semanal. El equipo ejecuta.'],
        ['ESTRATEGIA', 'ARS 550.000', 'Todo lo anterior + briefs creativos con copies + reunión mensual'],
        ['GESTIÓN TOTAL', 'ARS 750.000', 'Todo lo anterior + ejecución directa sobre campañas'],
    ],
    col_widths=[3.5, 3.5, 8.5]
)
doc.add_paragraph()

# Caja destacada ROI
p_roi = doc.add_paragraph()
p_roi.paragraph_format.space_before = Pt(8)
p_roi.paragraph_format.space_after = Pt(8)
r_roi = p_roi.add_run('💡  El plan se paga con recuperar el 8% de eficiencia en el primer mes. En los meses siguientes, el excedente es ganancia pura.')
r_roi.italic = True
r_roi.font.size = Pt(11)
r_roi.font.color.rgb = VIOLET
r_roi.font.name = 'Calibri'

# ── SECCIÓN 5 ─────────────────────────────────────────────────────────────────
add_heading(doc, '05. Por qué esto es diferente a una agencia', 1)
doc.add_paragraph()

add_table(doc,
    ['', 'Agencia tradicional', 'Este sistema'],
    [
        ['Velocidad de detección', 'Días / semanas', 'Minutos'],
        ['Reportes', 'Mensual o a pedido', 'Diario automático'],
        ['Decisiones basadas en', 'Criterio humano', 'Datos + IA'],
        ['Briefs de creativos', 'Extra / sin copies', 'Incluidos con copies listos'],
        ['Transparencia', 'Dashboard externo', 'Directo en tu Telegram'],
        ['Costo mensual', 'ARS 500k–1.500k', 'ARS 350k–750k'],
    ],
    col_widths=[5, 5.5, 5]
)

# ── SECCIÓN 6 ─────────────────────────────────────────────────────────────────
add_heading(doc, '06. Cómo empezamos', 1)
add_bullet(doc, 'Definimos el plan según las necesidades actuales del equipo', '1.  ')
add_bullet(doc, 'El sistema ya está corriendo sobre las campañas de IFPA — no hay setup ni período de adaptación', '2.  ')
add_bullet(doc, 'Primera auditoría del mes entregada dentro de las 48 horas', '3.  ')

doc.add_paragraph()
p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_final.paragraph_format.space_before = Pt(20)
r_final = p_final.add_run('Renzo Gasparini · renzogasparini85@gmail.com · Mayo 2026')
r_final.font.size = Pt(10)
r_final.font.color.rgb = GRAY
r_final.font.name = 'Calibri'

# Guardar
output = '/Users/renzogasparini/agente-meta-ads/Propuesta Sistema IA — IFPA Mayo 2026.docx'
doc.save(output)
print(f'Guardado: {output}')
